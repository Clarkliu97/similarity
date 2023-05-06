from django.http import HttpResponse
from rest_framework import generics, status
from rest_framework.response import Response
from .models import Task, File
from .tasks import similaritycheck
from .serializers import FileSerializer, TaskSerialiazer
from django.shortcuts import render
from rest_framework.views import APIView
from .models import Threshold
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from django.core.files.uploadedfile import InMemoryUploadedFile
from django.http.request import QueryDict
from docx import Document
from docx.opc.constants import CONTENT_TYPE as CT
from docx.oxml import parse_xml
from lxml import etree
import json
import io
import os
import requests
import csv

import openai
import docx
import matplotlib.pyplot as plt
from matplotlib.dates import date2num, DateFormatter
import datetime
import pandas as pd
import base64


output = []
# Create your views here.


def index(request):
    return render(request, "index.html")


############################################################################################
class DriveView(generics.ListAPIView, generics.GenericAPIView):
    '''
    This class receives http request about google drive
    '''
    authentication_classes = []
    permission_classes = []
    serializer_class = FileSerializer
    queryset = File.objects.all().order_by("-id")
    lookup_field = "id"

    def post(self, request):
        file = request.data.get("file", None)

        # Set up the OAuth2 authentication flow

        print("request: ", request.data)

        access_token = request.data.get('accessToken')
        auther = request.data.get('owners')[0].get('displayName')
        create_time = request.data.get('createdTime')
        modified_time = request.data.get('modifiedTime')
        modified_person = request.data.get('lastModifyingUser').get("displayName")
        creds = Credentials(token=access_token, scopes=['https://www.googleapis.com/auth/drive'])

        drive_service = build('drive', 'v3', credentials=creds)

        results = drive_service.files().list(
            fields="files(name,id,createdTime, modifiedTime, owners, lastModifyingUser)"
        ).execute()

        # Download a file from Google Drive
        file_id = request.data.get('fileid')
        request2 = drive_service.files().get_media(fileId=file_id)
        file_object = io.BytesIO()
        downloader = MediaIoBaseDownload(file_object, request2)
        done = False
        while done is False:
            stat, done = downloader.next_chunk()
            print("Download %d%%." % int(stat.progress() * 100))

        # Save the downloaded file to disk
        with open('tem.docx', 'wb') as f:
            f.write(file_object.getbuffer())

        def update_core_properties(docx_path, new_author, new_created_time, new_modified_time, last_modified_persion):

            """
            This function aims to change the metadata of docx but it seems that it doesn't work....
            """
            doc = Document(docx_path)
            core_properties = doc.core_properties
            core_properties.author = new_author
            core_properties.created = new_created_time
            core_properties.modified = new_modified_time
            core_properties.last_modified_by = last_modified_persion

            doc.save(docx_path)


            # Replace with the path to your .docx file

        input_docx_path = 'tem.docx'
        date_format = "%Y-%m-%dT%H:%M:%S.%fZ"
        # Update the core properties with the new author and created time

        new_author = auther
        new_created_time = datetime.strptime(create_time, date_format)
        new_modified_time = datetime.strptime(modified_time, date_format)
        last_modified_persion = modified_person

        update_core_properties(input_docx_path, new_author, new_created_time, new_modified_time, last_modified_persion)


        # Read the doc from disk and create a new InMemoryUploadedFile and pass it to serializer
        with open('tem.docx', 'rb') as f:
            file_content = f.read()

        file_object22 = io.BytesIO(file_content)

        uploaded_file = InMemoryUploadedFile(
            file_object22,
            field_name='file',
            name='tem.docx',
            content_type='application/msword',
            size=len(file_content),
            charset=None
        )

        query_dict = QueryDict(mutable=True)
        query_dict.update({'file': uploaded_file})
        temp = query_dict

        print("query_dict: ", query_dict)
        print("fileeee: ", query_dict.get("file"))

        if temp == "":
            return Response(
                {"file": "file is empty"}, status=status.HTTP_400_BAD_REQUEST
            )
        if temp is None:
            return Response(
                {"file": "file is none"}, status=status.HTTP_400_BAD_REQUEST
            )

        serializer = self.serializer_class(
            data=temp, context={"request": self.request}
        )
        print("is here")
        if serializer.is_valid(raise_exception=True):
            serializer.save()

            original_data = serializer.data

            openai.api_key = ("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")

            doc = docx.Document(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            file_text = '\n'.join(full_text)
            input_text = file_text

            if len(input_text) > 2000:
                input_text = input_text[:2000]

            attribute_dict = {
                "Personal": "Creates a sense of person/personality for the writer in the reader's mind. Typically written in a very informal style with first and/or second person and emotive language.",
                "Imaginative": "Uses the writer's creativity and imagination to entertain the reader. Uses techniques such as variation in sentence length, juxtaposition of different sentence lengths, careful control of structure and sequencing, to add to the overall effect by creating the desired atmosphere or conveying the required emotion.",
                "Persuasive": "Aims to convince readers about an idea, opinion or a course of action in order to achieve a desired outcome. Persuasive techniques chosen are influenced by the nature of the target audience; that is, the language (vocabulary, sentence structures, style/register), structure and sequencing of the piece are framed with the particular audience and purpose in mind",
                "Informative": "Generally uses facts, examples, explanations, analogies and sometimes statistical information, quotations and references as evidence. Chooses language, structure and sequence to make the message clear and unambiguous, so the sequencing of information is usually logical and predictable",
                "Evaluative": "Presents two or more important aspects of an issue or sides of an argument and discusses these rationally and objectively; using evidence to support the contrasting sides or alternatives. Uses objective style; appeals to reason not emotion; creation of an impression of balance and impartiality is essential.", }
            attribute_list = ["Geography, Agriculture and Environment", "Biology", "Mathematics",
                              "Culture, Arts and Music", "History", "Science",
                              "Education", "Health and Medicine", "English Language and Literature", "Technology",
                              "Physics", "Society", "Politics",
                              "Business and Economics", "Chemistry", "Legal studies", "Engineering", "Psychology",
                              "Architecture, Planning and Design",
                              "Journalism and Media", "Religious studies", "Other"]

            prompt1 = f"Which key in the dictionary {json.dumps(attribute_dict)} best represents the following text?\n\n{input_text}\n\nAnswer:"
            response1 = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt1,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
            )
            answer1 = response1.choices[0].text.strip()
            category_key = None
            for option in attribute_dict.keys():
                if option.lower() in answer1.lower():
                    category_key = option
                    break
            prompt2 = f"Which category in the list {json.dumps(attribute_list)} best represents the following text?\n\n{input_text}\n\nAnswer:"
            response2 = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt2,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
            )
            answer2 = response2.choices[0].text.strip()

            category_list = None
            for option in attribute_list:
                if option.lower() in answer2.lower():
                    category_list = option
                    break

            print(f"Category key: {category_key}")
            print(f"Category list: {category_list}")

            additional_data = {"category_key": category_key, "category_list": category_list}
            merged_data = {**original_data, **additional_data}
            output.append(merged_data)
            return Response(merged_data, status=status.HTTP_201_CREATED)

        print(serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def get(self, request, *args, **kwargs):
        return self.retrieve(request, *args, **kwargs)

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        response = HttpResponse(instance.file, content_type="text/txt")
        response["Content-Disposition"] = "attachment; filename={}".format(
            str(instance.file.name)
        )
        return response


class UploadFile(generics.ListAPIView, generics.GenericAPIView):
    authentication_classes = []
    permission_classes = []
    serializer_class = FileSerializer
    queryset = File.objects.all().order_by("-id")
    lookup_field = "id"



    def post(self, request):
        # print("file:", request)
        # print("---")

        print("file:", request.data)
        file = request.data.get("file", None)


        if file == "":
            return Response(
                {"file": "file is empty"}, status=status.HTTP_400_BAD_REQUEST
            )
        if file is None:
            return Response(
                {"file": "file is none"}, status=status.HTTP_400_BAD_REQUEST
            )

        serializer = self.serializer_class(
            data=request.data, context={"request": self.request}
        )
        if serializer.is_valid(raise_exception=True):
            serializer.save()
            original_data = serializer.data

            openai.api_key = ("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")

            doc = docx.Document(file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            file_text = '\n'.join(full_text)
            input_text = file_text

            if len(input_text) > 2000:
                input_text = input_text[:2000]

            attribute_dict = {
                "Personal": "Creates a sense of person/personality for the writer in the reader's mind. Typically written in a very informal style with first and/or second person and emotive language.",
                "Imaginative": "Uses the writer's creativity and imagination to entertain the reader. Uses techniques such as variation in sentence length, juxtaposition of different sentence lengths, careful control of structure and sequencing, to add to the overall effect by creating the desired atmosphere or conveying the required emotion.",
                "Persuasive": "Aims to convince readers about an idea, opinion or a course of action in order to achieve a desired outcome. Persuasive techniques chosen are influenced by the nature of the target audience; that is, the language (vocabulary, sentence structures, style/register), structure and sequencing of the piece are framed with the particular audience and purpose in mind",
                "Informative": "Generally uses facts, examples, explanations, analogies and sometimes statistical information, quotations and references as evidence. Chooses language, structure and sequence to make the message clear and unambiguous, so the sequencing of information is usually logical and predictable",
                "Evaluative": "Presents two or more important aspects of an issue or sides of an argument and discusses these rationally and objectively; using evidence to support the contrasting sides or alternatives. Uses objective style; appeals to reason not emotion; creation of an impression of balance and impartiality is essential.", }
            attribute_list = ["Geography, Agriculture and Environment", "Biology", "Mathematics",
                              "Culture, Arts and Music", "History", "Science",
                              "Education", "Health and Medicine", "English Language and Literature", "Technology",
                              "Physics", "Society", "Politics",
                              "Business and Economics", "Chemistry", "Legal studies", "Engineering", "Psychology",
                              "Architecture, Planning and Design",
                              "Journalism and Media", "Religious studies", "Other"]

            prompt1 = f"Which key in the dictionary {json.dumps(attribute_dict)} best represents the following text?\n\n{input_text}\n\nAnswer:"
            response1 = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt1,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
            )
            answer1 = response1.choices[0].text.strip()
            category_key = None
            for option in attribute_dict.keys():
                if option.lower() in answer1.lower():
                    category_key = option
                    break
            prompt2 = f"Which category in the list {json.dumps(attribute_list)} best represents the following text?\n\n{input_text}\n\nAnswer:"
            response2 = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt2,
                max_tokens=50,
                n=1,
                stop=None,
                temperature=0.5,
            )
            answer2 = response2.choices[0].text.strip()

            category_list = None
            for option in attribute_list:
                if option.lower() in answer2.lower():
                    category_list = option
                    break

            print(f"Category key: {category_key}")
            print(f"Category list: {category_list}")

            additional_data = {"category_key": category_key, "category_list": category_list}
            merged_data = {**original_data, **additional_data}
            output.append(merged_data)
            return Response(merged_data, status=status.HTTP_201_CREATED)
        print(serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def get(self, request, *args, **kwargs):
        return self.retrieve(request, *args, **kwargs)

    def retrieve(self, request, *args, **kwargs):
        instance = self.get_object()
        response = HttpResponse(instance.file, content_type="text/txt")
        response["Content-Disposition"] = "attachment; filename={}".format(
            str(instance.file.name)
        )
        return response


class TaskView(
    generics.RetrieveAPIView, generics.UpdateAPIView, generics.GenericAPIView
):
    authentication_classes = []
    permission_classes = []
    serializer_class = TaskSerialiazer
    queryset = Task.objects.all().order_by("-id")
    lookup_field = "id"

    def post(self, request):
        files = request.data.get("file_id")
        print("post,here:", files)
        authors = request.data.get("authors")
        if len(files) < 2:
            return Response(
                {"file": "Add more files"}, status=status.HTTP_400_BAD_REQUEST
            )
        if files is None or files == []:
            return Response(
                {"file": "Select atleast two file"}, status=status.HTTP_400_BAD_REQUEST
            )
        if authors is None or len(authors) == 0:
            return Response(
                {"author": "Select atleast one author"},
                status=status.HTTP_400_BAD_REQUEST,
            )

        # authors = File.objects.filter(id__in=files).values_list('author', flat=True)
        # authors = list(set(authors))
        # if len(objts) <= 1:
        #     return Response(
        #         {"file": "Select atleast two files for {}".format(authors)},
        #         status=status.HTTP_400_BAD_REQUEST,
        #     )

        # create and task object
        SimilarityCheck_obj = Task()
        SimilarityCheck_obj.authors = json.dumps(authors)
        SimilarityCheck_obj.save()

        # celery task
        # print("Authors:", authors)
        similaritycheck.delay(file=files, authors=authors, id=SimilarityCheck_obj.id)
        # file_list = get_all_legal_files(file=files, authors=authors, id=SimilarityCheck_obj.id)
        # print("file_list:\n", file_list)
        # Check every entry in output list and get the author's work
        result = remove_different_author(authors)
        result = remove_illegal_files(result)
        print("result:\n", result, "\n", len(result))
        result = error_checking(result)
        # for o in output:
        #     print("o:", o)
        plot_graph_by_date()

        write_to_html(authors)
        
        print("Files submitted sucsessfully", "task_id", SimilarityCheck_obj.id)
        return Response({"file": "Files submitted sucsessfully", "task_id": SimilarityCheck_obj.id},
                        status=status.HTTP_200_OK)

    def get(self, request, *args, **kwargs):
        print("get,here:", request)
        return self.retrieve(request, *args, **kwargs)

    # def patch(self, request, *args, **kwargs):
    #     return self.partial_update(request, *args, **kwargs)

    # def update(self, request, *args, **kwargs):
    #     partial = kwargs.pop("partial", False)
    #     instance = self.get_object()
    #     completed_file = instance.completed_file
    #     threshold_file = instance.threshold_file
    #     if completed_file < threshold_file:
    #         return Response(
    #             {"file": "upload more file "}, status=status.HTTP_400_BAD_REQUEST
    #         )
    #     serializer = self.get_serializer(instance, data=request.data, partial=partial)
    #     if serializer.is_valid():
    #         self.perform_update(serializer)
    #         return Response(serializer.data, status=status.HTTP_201_CREATED)
    #     else:
    #         return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class ConfigurationsView(APIView):
    def get(self, request):
        queryset = Threshold.get_solo()
        return Response({'threshold': queryset.min_files, "show_error": queryset.show_file_error},
                        status=status.HTTP_200_OK)


def remove_different_author(authors): 
    tmp = []
    global output
    for i in output: 
        auth = i["author"]
        for a in authors: 
            if a == auth: 
                # change all single quote to double quote
                tmp.append(i)
                break
    # print("output:\n", tmp, "\n", len(tmp))
    output = []
    return tmp


def remove_illegal_files(output): 
    tmp = []
    for o in output: 
        if o["is_error"] == False: 
            tmp.append(o)
    return tmp


def error_checking(output): 
    tot_nums = [] # total error numbers
    spe_nums = [] # spelling error numbers
    gra_nums = [] # grammar error numbers
    pun_nums = [] # punctuation error numbers
    syn_nums = [] # syntax error numbers
    sty_nums = [] # style error numbers
    oth_nums = [] # other error numbers

    list_of_id = []
    for i in output:
        id = i['id']
        list_of_id.append(id)
    print("list_of_id:\n", list_of_id)
    files = File.objects.filter(id__in=list_of_id, is_error=False)

    list_of_name = []
    for i in output:
        name = i['file_name']
        #remove directory
        name = name.split('/')[-1]
        list_of_name.append(name)

    list_of_date = []
    list_of_genre = []
    for i in output:
        genre = i['category_key']
        list_of_genre.append(genre)
    list_of_subject = []
    for i in output:
        subject = i['category_list']
        list_of_subject.append(subject)
    
    API_KEY = "917c943b-ed2f-41b1-a251-1303d17dc515" 
    url = "https://prevprod.gingersoftware.com/correction/v1/document?apiKey=" + API_KEY + "&lang=UK"
    headers = {"Content-Type": "text/plain"}

    category_dict = {
        **dict.fromkeys([47, 1000, 34, 2, 45], "Spelling"),
        **dict.fromkeys([29, 5, 51, 21, 40, 105, 38, 37, 20, 43, 36, 19, 35, 18, 16], "Grammar"),
        **dict.fromkeys([46, 50, 42, 100], "Punctuation"),
        **dict.fromkeys([44, 15, 48, 13, 30, 31, 12, 23, 39], "Syntax"),
        **dict.fromkeys([102, 103, 104, 49], "Style"),
    }

    for file in files: 
        content = file.text[0]
        content = content.encode('utf-8')
        errflag = file.text[1]
        # print("content:\n", content)
        date = file.created_at
        list_of_date.append(date)

        total_error_number, spelling_error_number, grammar_error_number, punctuation_error_number, syntax_error_number, style_error_number, other_error_number = 0, 0, 0, 0, 0, 0, 0
        if errflag == 1:
            tot_nums.append(total_error_number)
            spe_nums.append(spelling_error_number)
            gra_nums.append(grammar_error_number)
            pun_nums.append(punctuation_error_number)
            syn_nums.append(syntax_error_number)
            sty_nums.append(style_error_number)
            oth_nums.append(other_error_number)
            continue

        response = requests.post(url, headers=headers, data=content)
        if response.status_code == 200: 
            try: 
                result = json.loads(response.text)
                result = json.dumps(result, indent=2)
                with open('Single_Functions_by_Peicheng\ginger_error_tmp.json', 'w') as f:
                    f.write(result)
                    f.close()
                with open('Single_Functions_by_Peicheng\ginger_error_tmp.json', 'r') as f:
                    result = json.load(f)
                    f.close()
                # print("result:\n", result)

                if "GingerTheDocumentResult" in result: 
                    documentResult = result["GingerTheDocumentResult"]
                    errors = documentResult["Corrections"] 
                    for error in errors: 
                        # print(error)
                        categoryId = int(error["TopCategoryId"])
                        # print(categoryId)
                        total_error_number += 1
                        supId = category_dict.get(categoryId, "Other")
                        # print(supId)
                        if supId == "Spelling": 
                            spelling_error_number += 1
                        elif supId == "Grammar":
                            grammar_error_number += 1
                        elif supId == "Punctuation":
                            punctuation_error_number += 1
                        elif supId == "Syntax":
                            syntax_error_number += 1
                        elif supId == "Style":
                            style_error_number += 1
                        elif supId == "Other":
                            other_error_number += 1
                
            except json.decoder.JSONDecodeError as e:
                errString = "Failed to parse JSON response"
                print(f"Failed to parse JSON response: {e}")
                total_error_number, spelling_error_number, grammar_error_number, punctuation_error_number, syntax_error_number, style_error_number, other_error_number = errString, errString, errString, errString, errString, errString, errString
        else:
            print(f"Request failed with status code {response.status_code}.")
            errString = "Request failed with status code " + str(response.status_code) + ": " + str(response.text["Message"])
            total_error_number, spelling_error_number, grammar_error_number, punctuation_error_number, syntax_error_number, style_error_number, other_error_number = errString, errString, errString, errString, errString, errString, errString

        tot_nums.append(total_error_number)
        spe_nums.append(spelling_error_number)
        gra_nums.append(grammar_error_number)
        pun_nums.append(punctuation_error_number)
        syn_nums.append(syntax_error_number)
        sty_nums.append(style_error_number)
        oth_nums.append(other_error_number)

        print("Total error number\t Spelling error number\t Grammar error number\t Punctuation error number\t Syntax error number\t Style error number\t Other error number")
        print(total_error_number, "\t", spelling_error_number, "\t", grammar_error_number, "\t", punctuation_error_number, "\t", syntax_error_number, "\t", style_error_number, "\t", other_error_number)
    
    print(len(list_of_name), " ", len(list_of_date), " ", len(list_of_genre), " ", len(list_of_subject), " ", len(tot_nums), " ", len(spe_nums), " ", len(gra_nums), " ", len(pun_nums), " ", len(syn_nums), " ", len(sty_nums), " ", len(oth_nums))
    write_to_csv(list_of_name, list_of_date, list_of_genre, list_of_subject, tot_nums, spe_nums, gra_nums, pun_nums, syn_nums, sty_nums, oth_nums)
    # Add to output list
    for i in range(len(output)):
        new = {"total_error_number": tot_nums[i], "spelling_error_number": spe_nums[i], "grammar_error_number": gra_nums[i], "punctuation_error_number": pun_nums[i], "syntax_error_number": syn_nums[i], "style_error_number": sty_nums[i], "other_error_number": oth_nums[i]}
        merged = {**output[i], **new}
        output[i] = merged

    return output


def write_to_csv(file_list, list_of_date, list_of_genre, list_of_subject, total_error_numbers, spelling_error_numbers, grammar_error_numbers, punctuation_error_numbers, syntax_error_numbers, style_error_numbers, other_error_numbers):
    # Write to Ginger_Error_Stats.csv file
    with open('Output\Ginger_Error_Stats.csv', 'w', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(['File Name', 'Date', 'Genre', 'Subject', 'Total Error Number', 'Spelling Error Number', 'Grammar Error Number', 'Punctuation Error Number', 'Syntax Error Number', 'Style Error Number', 'Other Error Number'])
        for i in range(len(file_list)):
            file_name = os.path.basename(file_list[i]) 
            writer.writerow([file_name, list_of_date[i], list_of_genre[i], list_of_subject[i], total_error_numbers[i], spelling_error_numbers[i], grammar_error_numbers[i], punctuation_error_numbers[i], syntax_error_numbers[i], style_error_numbers[i], other_error_numbers[i]])
        total = "Total"
        writer.writerow([total, "", "", "", sum_list(total_error_numbers), sum_list(spelling_error_numbers), sum_list(grammar_error_numbers), sum_list(punctuation_error_numbers), sum_list(syntax_error_numbers), sum_list(style_error_numbers), sum_list(other_error_numbers)])
        f.close()


def sum_list(list) -> int:
    sum = 0
    for i in list:
        # if is intger
        if isinstance(i, int):
            sum += i
    return sum


def plot_graph_by_date(): 
    # Open the csv file
    with open('Output\Ginger_Error_Stats.csv', 'r') as f:
        # Read all lines
        reader = csv.reader(f)
        data = list(reader)
        # remove last line
        data.pop()
        # Get the date info (accurate to Month)
        date = []
        for i in range(1, len(data)): 
            year_month = data[i][1][0:7]
            if year_month not in date: 
                date.append(year_month)
        date.sort()
        print(date)

        # Create lists of 0s same length as date
        tot_nums = [0.] * len(date)
        spe_nums = [0.] * len(date)
        gra_nums = [0.] * len(date)
        pun_nums = [0.] * len(date)
        syn_nums = [0.] * len(date)
        sty_nums = [0.] * len(date)
        oth_nums = [0.] * len(date)
        fil_nums = [0.] * len(date)

        for i in range(len(date)):
            for j in range(len(data)): 
                if date[i] == data[j][1][0:7]: 
                    tot_nums[i] += int(data[j][4])
                    spe_nums[i] += int(data[j][5])
                    gra_nums[i] += int(data[j][6])
                    pun_nums[i] += int(data[j][7])
                    syn_nums[i] += int(data[j][8])
                    sty_nums[i] += int(data[j][9])
                    oth_nums[i] += int(data[j][10])
                    fil_nums[i] += 1
        
        tot_nums = [x / y for x, y in zip(tot_nums, fil_nums)]
        spe_nums = [x / y for x, y in zip(spe_nums, fil_nums)]
        gra_nums = [x / y for x, y in zip(gra_nums, fil_nums)]
        pun_nums = [x / y for x, y in zip(pun_nums, fil_nums)]
        syn_nums = [x / y for x, y in zip(syn_nums, fil_nums)]
        sty_nums = [x / y for x, y in zip(sty_nums, fil_nums)]
        oth_nums = [x / y for x, y in zip(oth_nums, fil_nums)]
        
        print(tot_nums)
        print(spe_nums)
        print(gra_nums)
        print(pun_nums)
        print(syn_nums)
        print(sty_nums)
        print(oth_nums)
        print(fil_nums)

    for i in range(len(date)):
        date[i] = datetime.datetime.strptime(date[i], '%Y-%m')
        date[i] = date2num(date[i])

    plt.figure(figsize=(16, 8))
    plt.plot_date(date, tot_nums, label='Total', linestyle='solid', color='black')
    plt.plot_date(date, spe_nums, label='Spelling', linestyle='solid', color='red')
    plt.plot_date(date, gra_nums, label='Grammar', linestyle='solid', color='orange')
    plt.plot_date(date, pun_nums, label='Punctuation', linestyle='solid', color='yellow')
    plt.plot_date(date, syn_nums, label='Synonym', linestyle='solid', color='green')
    plt.plot_date(date, sty_nums, label='Style', linestyle='solid', color='blue')
    plt.plot_date(date, oth_nums, label='Other', linestyle='solid', color='purple')
    # Make x axis format as date
    plt.gca().xaxis.set_major_formatter(DateFormatter('%Y-%m'))
    plt.xlabel('Date')
    plt.ylabel('Average Number of Errors per Document')
    plt.legend()
    # plt.show()
    plt.savefig('Output\Ginger_Error_Stats.png')


def write_to_html(authors): 
    author = "Author: "
    for i in authors:
        author += i + ", "
    author = author[:-2]

    author_tag = "<p>Author: " + author + "</p>"

    table = pd.read_csv('Output\Ginger_Error_Stats.csv')
    # table.to_html('Output\Ginger_Error_Stats.html')
    table_html = table.to_html()

    img = base64.b64encode(open('Output\Ginger_Error_Stats.png', 'rb').read()).decode('utf-8')
    img_tag = '<img src="data:image/png;base64,{0}">'.format(img)
    # print(img_tag)

    html = author_tag + table_html + img_tag
    with open('Output\Ginger_Error_Stats.html', 'w') as f:
        f.write(html)
        f.close()
